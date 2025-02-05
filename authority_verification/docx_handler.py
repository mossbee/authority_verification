import docx
import json
from . import utils, config, doc_utils

class DocxHandler:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = None
        self.paragraphs_index = {}

    def read_docx(self):
        """
            Reads the .docx file and loads it into the document attribute.
        """
        self.document = docx.Document(self.file_path)
    
    def index_document(self):
        """
            Indexes the paragraphs in the document:
                ["chapter_index.article_index.clause_index.point_index"]
        """
        if self.document is None:
            raise ValueError("Document not loaded. Call read_docx() first.")
        
        document_text = []
        document_len = len(self.document.paragraphs)
        for i in range(0, document_len):
            if self.document.paragraphs[i].text != '\xa0' and self.document.paragraphs[i].text != '' and not self.document.paragraphs[i].text.startswith('Ví dụ'):
                document_text.append(self.document.paragraphs[i].text)
        document_len = len(document_text)
        document_data = [
            {
                "index" : "0.0.0.0",
                "text" : document_text[0]
            }
        ]

        for i in range(1, document_len):
            cur_element = document_text[i]
            prev_element_id = document_data[i - 1]["index"].split(".")
            cur_element_text_splitted = cur_element.split()
            if cur_element_text_splitted[0] == "Chương":
                # cur_element: Chương 1 lorem ipsum
                if cur_element_text_splitted[1].isdigit():
                    cur_element_ch_id = cur_element_text_splitted
                # cur_element: Chương 1. lorem ipsum or Chương 1: lorem ipsum
                elif not cur_element_text_splitted[1].isdigit() and cur_element_text_splitted[1][:-1].isdigit():
                    cur_element_ch_id = cur_element_text_splitted[1][:-1]
                # cur_element: Chương I lorem ipsum
                elif doc_utils.is_roman(cur_element_text_splitted[1]):
                    cur_element_ch_id = doc_utils.roman_to_int(cur_element_text_splitted[1])
                # cur_element: Chương I. lorem ipsum or Chương I: lorem ipsum
                elif not cur_element_text_splitted[1].isalpha() and doc_utils.is_roman(cur_element_text_splitted[1][:-1]):
                    cur_element_ch_id = doc_utils.roman_to_int(cur_element_text_splitted[1][:-1])
                document_data.append(
                    {
                        "index" : f"{cur_element_ch_id}.0.0.0",
                        "text" : cur_element
                    }
                )
            # cur_element: Điều lorem ipsum
            elif cur_element_text_splitted[0] == "Điều":
                # cur_element: Điều 1 lorem ipsum
                if cur_element_text_splitted[1].isdigit():
                    cur_element_ar_id = cur_element_text_splitted[1]
                # cur_element: Điều 1. lorem ipsum or Điều 1: lorem ipsum
                elif not cur_element_text_splitted[1].isdigit() and cur_element_text_splitted[1][:-1].isdigit():
                    cur_element_ar_id = cur_element_text_splitted[1][:-1]
                document_data.append(
                    {
                        "index" : f"{prev_element_id[0]}.{cur_element_ar_id}.0.0",
                        "text" : cur_element
                    }
                )
            # cur_element: 1. lorem ipsum or 1) lorem ipsum
            elif not cur_element_text_splitted[0].isdigit() and cur_element_text_splitted[0][:-1].isdigit():
                cur_element_cl_id = cur_element_text_splitted[0][:-1]
                document_data.append(
                    {
                        "index" : f"{prev_element_id[0]}.{prev_element_id[1]}.{cur_element_cl_id}.0",
                        "text" : cur_element
                    }
                )
            # cur_element: a. lorem ipsum or a) lorem ipsum
            elif not cur_element_text_splitted[0].isalpha() and cur_element_text_splitted[0][:-1].isalpha():
                cur_element_pt_id = cur_element_text_splitted[0][:-1]
                document_data.append(
                    {
                        "index" : f"{prev_element_id[0]}.{prev_element_id[1]}.{prev_element_id[2]}.{cur_element_pt_id}",
                        "text" : cur_element
                    }
                )
            # cur_element: lorem ipsum
            else:
                document_data.append(
                    {
                        "index" : f"{prev_element_id[0]}.{prev_element_id[1]}.{prev_element_id[2]}.{prev_element_id[3]}",
                        "text" : cur_element
                    }
                )
        # merge paragraphs with the same index
        for i in range(1, document_len):
            if document_data[i]["index"] == document_data[i - 1]["index"]:
                document_data[i - 1]["text"] += " " + document_data[i]["text"]
                document_data[i]["text"] = ""
        # remove empty paragraphs
        document_data = [para for para in document_data if para["text"] != ""]
        index_dict = {}
        for para in document_data:
            if para["index"] not in index_dict:
                index_dict[para["index"]] = para["text"]
            else:
                index_dict[para["index"]] += " " + para["text"]

        self.paragraphs_index = index_dict

    def save_indexed_paragraphs_to_json(self, output_file_path = config.OUTPUT_PATH):
        """
            Saves the indexed paragraphs to a .json file.
        """
        if not self.paragraphs_index:
            raise ValueError("No paragraphs indexed. Please index the paragraphs first.")
        
        with open(output_file_path + self.file_path.split("/")[-1][:-5] + "_indexed.json", 'w', encoding='utf-8') as json_file:
            json.dump(self.paragraphs_index, json_file, ensure_ascii=False, indent=4)