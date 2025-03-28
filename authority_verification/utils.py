import os
import re
import json
from tqdm import tqdm
from docx import Document
from typing import Dict, List
from unidecode import unidecode
from docx.enum.text import WD_COLOR_INDEX
from . import config, doc_utils, docx_handler

def get_name_from_path(file_path:str):
	"""
		Get the name of the file from the file path.

		Parameters:
			file_path (str): The path of the file.
		
		Returns:
			str: The name of the file.

		Example:
			>>> get_name_from_path('data/legal_cases/65/219_2013_TT-BTC_220761.docx')
			'219_2013_TT-BTC_220761'
	"""
	normalized_path = file_path.replace('\\', '/')
	return normalized_path.split("/")[-1][:-5]

def json_dumper(input_dict: Dict, file_path: str, purpose: str):
	"""
		Write a dictionary to a JSON file.

		Parameters:
			input_dict (Dict): The dictionary to be written to the JSON file.
			file_path (str): The path of the file.
			purpose (str): The purpose of the JSON file.
		
		Returns:
			None
	"""
	with open(config.OUTPUT_PATH + get_name_from_path(file_path) + purpose, 'w', encoding="utf-8") as json_file:
		json.dump(input_dict, json_file, indent = 4, ensure_ascii = False)

def load_json(file_path: str) -> Dict:
	"""
		Load a JSON file.

		Parameters:
			file_path (str): The path of the file.
		
		Returns:
			Dict: The dictionary loaded from the JSON file.
	"""
	with open(file_path, 'r', encoding='utf-8') as f:
		data = json.load(f)
	return data

def get_agencies_list():
	"""
		Get the list of agencies from the JSON file.

		Parameters:
			None
		
		Returns:
			list: The list of agencies.
	"""
	with open(config.VIETNAM_AGENCIES_LIST_PATH, 'r', encoding='utf-8') as f:
		vietnam_agencies = json.load(f)
	return vietnam_agencies["vietnam_agencies"]

def highlight_agencies(file: str):
	"""
		Highlight the agencies in a document.

		Parameters:
			file (str): The path of the document.
		
		Returns:
			None
	"""
	doc = Document(file)
	vietnam_agencies = get_agencies_list()
	# Iterate through paragraphs and runs to find and highlight the agencies
	for agency in vietnam_agencies:
		for paragraph in doc.paragraphs:
			for run in paragraph.runs:
				original_text = run.text
				normalized_text = unidecode(original_text).lower()
				if unidecode(agency).lower() in normalized_text:
					start = normalized_text.find(unidecode(agency).lower())
					end = start + len(unidecode(agency).lower())
					run.text = original_text[:start] + original_text[start:end] + original_text[end:]
					run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Highlight with yellow color
	
	# Save the modified document with the new name
	new_file_name = file.replace('.docx', '_highlighted.docx')
	doc.save(new_file_name)

def index_one_documents(document_path: str, save_dir = True):
	dochandler = docx_handler.DocxHandler(document_path)
	dochandler.read_docx()
	dochandler.index_document()

	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(document_path) + '_indexed.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(dochandler.paragraphs_index, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(document_path) + '_indexed.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(dochandler.paragraphs_index, indent = 4, ensure_ascii = False))
	
	return dochandler.paragraphs_index

def index_only_articles_name(document_path: str, save_dir: None):
	articles_name = {}
	doc_dict = index_one_documents(document_path, save_dir)
	for key, value in doc_dict.items():
		key_splitted = key.split('.')
		if key_splitted[2] == '0' and key_splitted[3] == '0' and key_splitted[1] != '0':
			articles_name[key] = value
		
	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(document_path) + '_articles_names.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(articles_name, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(document_path) + '_articles_names.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(articles_name, indent = 4, ensure_ascii = False))

def index_full_articles(document_path: str, save_dir: None):
	articles = {}
	doc_dict = index_one_documents(document_path, save_dir)
	doc_keys = list(doc_dict.keys())
	doc_len = len(doc_dict)
	i = 0
	while i < doc_len:
		key = doc_keys[i]
		key_splitted = key.split('.')
		if key_splitted[2] == '0' and key_splitted[3] == '0' and key_splitted[1] != '0':
			article_content = doc_dict[key] + ' '
			# concat all paragraphs that belong to an article
			while i < doc_len - 1 and doc_keys[i + 1].split('.')[1] == key_splitted[1]:
				article_content += doc_dict[doc_keys[i + 1]] + ' '
				i += 1
			
			articles[key] = article_content
		i += 1
	
	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(document_path) + '_articles_full.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(articles, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(document_path) + '_articles_full.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(articles, indent = 4, ensure_ascii = False))

def index_only_clauses_name(document_path: str, save_dir: None):
	clauses_name = {}
	doc_dict = index_one_documents(document_path, save_dir)
	for key, value in doc_dict.items():
		key_splitted = key.split('.')
		if key_splitted[1] != '0' and key_splitted[2] != '0' and key_splitted[3] == '0':
			clauses_name[key] = doc_dict['.'.join([key_splitted[0], key_splitted[1], '0', '0'])] + ' ' + value
	
	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(document_path) + '_clauses_names.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(clauses_name, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(document_path) + '_clauses_names.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(clauses_name, indent = 4, ensure_ascii = False))

def index_full_clauses(document_path: str, save_dir: None):
	clauses = {}
	doc_dict = index_one_documents(document_path, save_dir)
	doc_keys = list(doc_dict.keys())
	doc_len = len(doc_dict)
	i = 0
	while i < doc_len:
		key = doc_keys[i]
		key_splitted = key.split('.')
		if key_splitted[1] != '0' and key_splitted[2] != '0' and key_splitted[3] == '0':
			clause_content = doc_dict['.'.join([key_splitted[0], key_splitted[1], '0', '0'])] + ' ' + doc_dict[key] + ' '
			# concat all paragraphs that belong to a clause
			while i < doc_len - 1 and doc_keys[i + 1].split('.')[1] == key_splitted[1] and doc_keys[i + 1].split('.')[2] == key_splitted[2]:
				clause_content += doc_dict[doc_keys[i + 1]] + ' '
				i += 1

			clauses[key] = clause_content
		i += 1
	
	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(document_path) + '_clauses_full.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(clauses, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(document_path) + '_clauses_full.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(clauses, indent = 4, ensure_ascii = False))

def index_points(document_path: str, save_dir: None):
	points_name = {}
	doc_dict = index_one_documents(document_path, save_dir)
	for key, value in doc_dict.items():
		key_splitted = key.split('.')
		if key_splitted[1] != '0' and key_splitted[2] != '0' and key_splitted[3] != '0':
			points_name[key] = doc_dict['.'.join([key_splitted[0], key_splitted[1], '0', '0'])] + ' ' + doc_dict['.'.join([key_splitted[0], key_splitted[1], key_splitted[2], '0'])] + ' ' + value

	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(document_path) + '_points.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(points_name, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(document_path) + '_points.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(points_name, indent = 4, ensure_ascii = False))

def find_ref_one_document(file_path: str):
	doc = docx_handler.DocxHandler(file_path)
	doc.read_docx()
	doc.index_document()
	doc.save_indexed_paragraphs_to_json()
	doc_dict = doc.paragraphs_index
	output = {}
	for key, value in doc_dict.items():
		ref_list = doc_utils.extract_reference_from_txt(value, key)
		if ref_list:
			output[key] = ref_list
		else:
			output[key] = []
	with open(config.OUTPUT_PATH + get_name_from_path(file_path) + '_ref.json', 'w', encoding="utf-8") as json_file:
		json_file.write(json.dumps(output, indent = 4, ensure_ascii = False))

def find_ref_all_legal_docs():
	for root, dirs, files in os.walk(config.LEGAL_DOCS_PATH):
		for file in tqdm(files):
			file_path = os.path.join(config.LEGAL_DOCS_PATH, file)
			print(file_path)
			find_ref_one_document(file_path)
			print("Done processing " + file_path)

def remove_unwanted_from_doc(file_path: str, save_dir: None):
	output = {}
	doc = docx_handler.DocxHandler(file_path)
	doc.read_docx()
	doc.index_document()
	doc.save_indexed_paragraphs_to_json()
	doc_dict = doc.paragraphs_index
	agencies_list = load_json(config.VIETNAM_AGENCIES_LIST_PATH)["vietnam_agencies"]
	for key, value in doc_dict.items():
		for agency in agencies_list:
			if agency in value:
				output[key] = value
	
	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(file_path) + '_filtered.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(output, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(file_path) + '_filtered.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(output, indent = 4, ensure_ascii = False))

	return output

def extract_jurisdiction_doc(file_path: str, save_dir: None):
	filtered_docs = remove_unwanted_from_doc(file_path, save_dir)
	output = {}
	for key, value in filtered_docs.items():
		extracted = doc_utils.juris_extract(value)
		if extracted:
			output[key] = extracted
	
	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(file_path) + '_jurisdiction.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(output, indent = 4, ensure_ascii = False))
	
	else:
		with open(save_dir + "/" + get_name_from_path(file_path) + '_jurisdiction.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(output, indent = 4, ensure_ascii = False))

	return output

def remove_unwanted_docs_all():
	for root, dirs, files in os.walk(config.LEGAL_DOCS_PATH):
		for file in tqdm(files):
			file_path = os.path.join(config.LEGAL_DOCS_PATH, file)
			print(file_path)
			remove_unwanted_from_doc(file_path)
			print("Done processing " + file_path)

def find_match_key(input_dict: Dict, input_key: str):
	output = []
	input_key_list = input_key.split('.')
	if input_key_list[3] != '0':
		for key in input_dict.keys():
			if key.split('.')[1] == input_key_list[1] and key.split('.')[2] == input_key_list[2] and key.split('.')[3] == input_key_list[3]:
				output.append(key)
				break
	if input_key_list[3] == '0' and input_key_list[2] != '0':
		for key in input_dict.keys():
			if key.split('.')[1] == input_key_list[1] and key.split('.')[2] == input_key_list[2]:
				output.append(key)
	if input_key_list[3] == '0' and input_key_list[2] == '0' and input_key_list[1] != '0':
		for key in input_dict.keys():
			if key.split('.')[1] == input_key_list[1]:
				output.append(key)
	return output

def find_belongings(input_dict: Dict, input_key: str):
	output = [input_dict[input_key]]
	first_zero_index = 4
	input_key_list = input_key.split('.')
	for i in range(len(input_key_list)):
		if input_key_list[i] != '0':
			for j in range(i + 1, len(input_key_list)):
				if input_key_list[j] == '0':
					first_zero_index = j
					break

	# print(first_zero_index)
	
	if first_zero_index == 4:
		return output
	elif first_zero_index == 3:
		text_placeholder = []
		for key in input_dict.keys():
			key_splitted = key.split('.')
			if key_splitted[1] == input_key_list[1] and key_splitted[2] == input_key_list[2]:
				text_placeholder.append(input_dict[input_key] + ' ' + input_dict[key])
		output += text_placeholder
		return output
	elif first_zero_index == 2:
		text_placeholder = []
		for key in input_dict.keys():
			key_splitted = key.split('.')
			if key_splitted[3] != '0' and key_splitted[1] == input_key_list[1]:
				text_placeholder.append(input_dict[input_key] + ' ' + input_dict['.'.join(key_splitted[:3]) + '.0'] + ' ' + input_dict[key])
			elif key_splitted[3] == '0' and key_splitted[1] == input_key_list[1]:
				text_placeholder.append(input_dict[input_key] + ' ' + input_dict[key])
		output += text_placeholder
		return output

def jurisdict_augmentation(file_path: str, save_dir: None):
	doc_filtered = {}
	doc_juris = {}
	doc_ref = {}
	doc = docx_handler.DocxHandler(file_path)
	doc.read_docx()
	doc.index_document()
	if save_dir is None:
		doc.save_indexed_paragraphs_to_json()
	else:
		with open(save_dir + "/" + get_name_from_path(file_path) + '_indexed.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(doc.paragraphs_index, indent = 4, ensure_ascii = False))
	doc_dict = doc.paragraphs_index
	agencies_list = load_json(config.VIETNAM_AGENCIES_LIST_PATH)["vietnam_agencies"]
	for key, value in doc_dict.items():
		for agency in agencies_list:
			if agency.lower() in value.lower():
				doc_filtered[key] = value
				break
	
	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(file_path) + '_filtered.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(doc_filtered, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(file_path) + '_filtered.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(doc_filtered, indent = 4, ensure_ascii = False))

	for key, value in doc_filtered.items():
		extracted = doc_utils.juris_extract(value)
		if extracted:
			below_content = []
			for i in range(len(extracted["content"])):
				extracted["content"][i] = extracted["content"][i].strip()
				if len(extracted["content"][i].split()) > 1 and extracted["content"][i].split()[-1][-1] == ':':
					# print(extracted["content"][i].split()[-1])
					# key_splitted = key.split('.')
					# # first_zero_index = key_splitted.index('0')
					# # find the first non zero index
					# for j in range(len(key_splitted)):
					# 	if key_splitted[j]!= '0':
					# 		first_zero_index = j
					# 		# now find the first zero index since this position
					# 		for k in range(first_zero_index, len(key_splitted)):
					# 			if key_splitted[k] == '0':
					# 				first_zero_index = k
					# 				break

					# for ori_key in doc_dict.keys():
					# 	if ori_key.split('.')[:first_zero_index] == key_splitted[:first_zero_index]:
					# 		print(ori_key.split('.')[:first_zero_index])
					# 		if ori_key.split('.')[first_zero_index] != '0' and ori_key.split('.')[first_zero_index] > key_splitted[first_zero_index]:
					# 			below_content.append(extracted["content"][i] + ' ' + doc_dict[ori_key])
					below_content += find_belongings(doc_dict, key)
					# print(below_content)
				
			extracted["content"] += below_content

			add_extract = []
			for i in range(len(extracted["content"])):
				ref_list = doc_utils.extract_reference_from_txt(extracted["content"][i], key)
				if ref_list:
					for reference in ref_list:
						text_holder = ' '
						matched_keys = find_match_key(doc_dict, reference)
						if matched_keys:
							for matched_key in matched_keys:
								text_holder += doc_dict[matched_key] + ' '
							add_extract.append(extracted["content"][i] + ' ' + text_holder)
			
			extracted["content"] += add_extract
			doc_juris[key] = extracted
	
	if save_dir is None:
		with open(config.OUTPUT_PATH + get_name_from_path(file_path) + '_jurisdiction.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(doc_juris, indent = 4, ensure_ascii = False))
	else:
		with open(save_dir + "/" + get_name_from_path(file_path) + '_jurisdiction.json', 'w', encoding="utf-8") as json_file:
			json_file.write(json.dumps(doc_juris, indent = 4, ensure_ascii = False))

	# for key, value in doc_juris.items():
	# 	contents = value["content"]
	# 	for content in contents:
	# 		ref_list = doc_utils.extract_reference_from_txt(content, key)
	# 		if ref_list:
	# 			doc_ref[key] = ref_list
	# 			for reference in ref_list:
	# 				matched_keys = find_match_key(doc_dict, reference)
	# 				if matched_keys:
	# 					for matched_key in matched_keys:
	# 						doc_filtered[key] += ' ' + doc_dict[matched_key]
	# 		else:
	# 			doc_ref[key] = []
	
	# if is_save_ref:
	# 	with open(config.OUTPUT_PATH + get_name_from_path(file_path) + '_ref.json', 'w', encoding="utf-8") as json_file:
	# 		json_file.write(json.dumps(doc_ref, indent = 4, ensure_ascii = False))
	
	# with open(config.OUTPUT_PATH + get_name_from_path(file_path) + '_augmented.json', 'w', encoding="utf-8") as json_file:
	# 	json_file.write(json.dumps(doc_filtered, indent = 4, ensure_ascii = False))

def table_extraction(document_path: str):
    dochandler = docx_handler.DocxHandler(document_path)
    dochandler.read_docx()
    document = dochandler.document
    table = document.tables[6]
    # print(table)
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data

def taxonomy_construction(document_path: str):
    data = {"Ngành, nghề đầu tư kinh doanh có điều kiện" : []}
    table_datas = table_extraction(document_path)
    for table_data in table_datas:
        data["Ngành, nghề đầu tư kinh doanh có điều kiện"].append(table_data["Ngành, nghề"])
    with open(config.KNOWLEDGE_GRAPH_PATH + 'knowledge_graph_conditional.json', 'w', encoding="utf-8") as json_file:
        json.dump(data, json_file, indent = 4, ensure_ascii = False)

if __name__ == "__main__":
	# find_ref_one_document(config.PURSUANT_DOCUMENT_PATH)
	find_ref_all_legal_docs()